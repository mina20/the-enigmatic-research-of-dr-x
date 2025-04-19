import os
import pandas as pd
from docx import Document
import fitz  
import glob
# from pprint import pprint
# from helper import get_llama_documents
class FileTextExtractor:
    def __init__(self, directory):
        self.directory = directory
        self.supported_extensions = [".docx", ".pdf", ".xls", ".xlsx", ".xlsm", ".csv"]
        self.all_data = []

    def extract_text_from_docx(self, file_path):
        data = []
        doc = Document(file_path)
        text_chunks = [para.text for para in doc.paragraphs if para.text.strip()]
        data.append({
            "file_name": os.path.basename(file_path),
            "file_type": "docx",
            "page_or_sheet": None,
            "content": "\n".join(text_chunks)
        })

        for table in doc.tables:
            table_text = []
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                table_text.append(" | ".join(row_text))
            data.append({
                "file_name": os.path.basename(file_path),
                "file_type": "docx",
                "page_or_sheet": "table",
                "content": "\n".join(table_text)
            })

        return data

    def extract_text_from_pdf(self, file_path):
        data = []
        with fitz.open(file_path) as doc:
            for page_num, page in enumerate(doc, 1):
                text = page.get_text().strip()
                if text:
                    data.append({
                        "file_name": os.path.basename(file_path),
                        "file_type": "pdf",
                        "page_or_sheet": f"page_{page_num}",
                        "content": text
                    })
        return data

    def extract_text_from_excel(self, file_path):
        data = []
        xls = pd.ExcelFile(file_path)
        for sheet_name in xls.sheet_names:
            df = xls.parse(sheet_name)
            rows = [" | ".join(map(str, row)) for row in df.itertuples(index=False)]
            if rows:
                data.append({
                    "file_name": os.path.basename(file_path),
                    "file_type": "excel",
                    "page_or_sheet": sheet_name,
                    "content": "\n".join(rows)
                })
        return data

    def extract_text_from_csv(self, file_path):
        data = []
        df = pd.read_csv(file_path)
        rows = [" | ".join(map(str, row)) for row in df.itertuples(index=False)]
        if rows:
            data.append({
                "file_name": os.path.basename(file_path),
                "file_type": "csv",
                "page_or_sheet": None,
                "content": "\n".join(rows)
            })
        return data

    def extract_text_from_file(self, file_path):
        ext = os.path.splitext(file_path)[1].lower()
        if ext == ".docx":
            return self.extract_text_from_docx(file_path)
        elif ext == ".pdf":
            return self.extract_text_from_pdf(file_path)
        elif ext in [".xls", ".xlsx", ".xlsm"]:
            return self.extract_text_from_excel(file_path)
        elif ext == ".csv":
            return self.extract_text_from_csv(file_path)
        else:
            raise ValueError(f"Unsupported file format: {ext}")

    def extract_all(self):
        for file in glob.glob(os.path.join(self.directory, "*")):
            try:
                if os.path.splitext(file)[1].lower() in self.supported_extensions:
                    file_data = self.extract_text_from_file(file)
                    self.all_data.extend(file_data)
            except Exception as e:
                print(f"Error reading {file}: {e}")
        return self.all_data
# if __name__ == "__main__":
#     directory = "../data/Dr_x_files"  
#     extractor = FileTextExtractor(directory)
#     all_data = extractor.extract_all()
    # Convert to LlamaIndex Document format
    # llama_docs = get_llama_documents(all_data)
    # pprint(all_data[0]) 
    # pprint("Extracted data:")
    # pprint(llama_docs[0])

